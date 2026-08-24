from pathlib import Path


def extract_pdf_text(path: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pypdf is not installed.") from exc

    try:
        reader = PdfReader(path)
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:
                raise ValueError("This PDF is password-protected and cannot be read.") from exc

        pages = []
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception:
                pages.append("")
        return "\n".join(pages).strip()
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError("The PDF appears to be corrupted or unreadable.") from exc


def extract_docx_text(path: str) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed.") from exc

    try:
        document = Document(path)
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return "\n".join(paragraphs).strip()
    except Exception as exc:
        raise ValueError("The DOCX file appears to be invalid or unreadable.") from exc


def extract_txt_text(path: str) -> str:
    encodings = ("utf-8", "utf-8-sig", "cp1252", "latin-1")

    for encoding in encodings:
        try:
            return Path(path).read_text(encoding=encoding).strip()
        except UnicodeDecodeError:
            continue
        except OSError as exc:
            raise ValueError("The text file could not be read.") from exc

    raise ValueError("The text file uses an unsupported character encoding.")


def extract_document_text(path: str) -> str:
    extension = Path(path).suffix.lower()

    if extension == ".pdf":
        return extract_pdf_text(path)
    if extension == ".docx":
        return extract_docx_text(path)
    if extension == ".txt":
        return extract_txt_text(path)

    raise ValueError("Unsupported document format.")
